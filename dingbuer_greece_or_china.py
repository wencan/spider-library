#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# wencan
# 2017-09-29 09:02

import requests
import lxml
import lxml.html
import lxml.etree
import parse
import chardet
import time
import urllib.parse
import collections
import io
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


home = "http://www.ikexue.org/specials/dingbuer_greece_or_china"
encoding = "UTF-8"


session = requests.Session()
session.headers.update({
    "User-Agent": "Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; WOW64; Trident/5.0)",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Encoding": "gzip, deflate, sdch, br",
    "Accept-Language": "zh-CN,zh;q=0.8,en;q=0.6,zh-TW;q=0.4"
})


def get(url, encoding='utf-8', sleep_seconds=0):
    print("GET:", url)

    try:
        r = session.get(url, allow_redirects=True)
    except requests.exceptions.ConnectionError:
        if sleep_seconds == 0:
            sleep_seconds = 10
        time.sleep(sleep_seconds)

        return get(url, encoding, sleep_seconds + sleep_seconds)

    if r.status_code == 200:
        pass
    elif r.status_code == 404:
        raise LookupError("not found %s" % url)
    else:
        raise IOError('request "%s" got status: %d' % (url, r.status_code))

    if encoding is None:
        return r.content

    try:
        text = r.content.decode(encoding)
    except UnicodeDecodeError:
        encoding = chardet.detect(r.content)["encoding"]
        text = r.content.decode(encoding)

    return text

Chapter = collections.namedtuple("Chapter", ["title", "sections"])
Section = collections.namedtuple("Section", ["title", "url"])
Image = collections.namedtuple("Image", ["url", "width", "height"])
Paragraph = collections.namedtuple("Paragraph", ["text", "img", "table"])


def get_preface_and_chapters(home_url, encoding="UTF-8"):
    text = get(home_url, encoding)
    doc = lxml.html.fromstring(text)

    body = doc.xpath("/html/body")[0]
    preface = body.xpath("//div[@class=\"entry-content\"]/div[@class=\"single-content\"]/p[@class=\"zt\"]")[0]
    preface = preface.text

    chapters = []
    divs = body.xpath("//div[@class=\"zt_sticky\"]")
    for div in divs:
        chapter_title = div.xpath(".//h2")[0]
        chapter_title = chapter_title.text

        sections = []

        hlinks = div.getnext().getnext()
        links = hlinks.xpath(".//li/a")
        for link in links:
            section_title = link.text

            url = link.attrib["href"]
            url = urllib.parse.urljoin(home_url, url)

            sections.append(Section(section_title, url))

        hlinks = hlinks.getnext()
        links = hlinks.xpath(".//li/a")
        for link in links:
            section_title = link.text

            url = link.attrib["href"]
            url = urllib.parse.urljoin(home_url, url)

            sections.append(Section(section_title, url))

        sections = sorted(sections, key=lambda section: int(parse.parse("{}、{}", section.title)[0]))
        chapters.append(Chapter(chapter_title, sections))

    return preface, chapters


def get_paragraphs(section_url, encoding="UTF-8"):
    text = get(section_url, encoding)
    doc = lxml.html.fromstring(text)

    body = doc.xpath("/html/body")[0]
    lines = body.xpath(".//div[@class=\"entry-content\"]/div[@class=\"single-content\"]/*")

    paragraphs = []

    for line in lines:
        if line.text and len(line.text.lstrip().rstrip()):
            if line.tag == "p" and line.xpath("@style=\"text-align: center\""):
                continue

            p = Paragraph(text=line.text, img=None, table=None)
            paragraphs.append(p)
        elif line.xpath(".//img[@src]"):
            img = line.xpath(".//img[@src]")[0]
            url = img.attrib["src"]
            url = urllib.parse.urljoin(section_url, url)
            text = None

            width, height = None, None
            if "width" in img.attrib:
                width = int(img.attrib["width"])
            if "height" in img.attrib:
                height = int(img.attrib["height"])

            img = Image(url=url, width=width, height=height)

            captions = line.xpath(".//p[@class=\"wp-caption-text\"]")
            if captions:
                caption = captions[0]
                text = caption.text
            else:
                line = line.getnext()
                if line.tag == "p":
                    if line.xpath("@style=\"text-align: center\""):
                        if line.text:
                            text = line.text
                        else:
                            ems = line.xpath("em")
                            if ems:
                                em = ems[0]
                                if em.text:
                                    text = em.text

            p = Paragraph(img=img, text=text, table=None)
            paragraphs.append(p)
        elif line.xpath(".//table"):
            table = line.xpath(".//table")[0]

            rows = []
            trs = table.xpath(".//tr")
            for tr in trs:
                row = []
                tds = tr.xpath("td")
                for td in tds:
                    row.append(td.text or "")
                rows.append(row)

            p = Paragraph(img=None, text=None, table=rows)
            paragraphs.append(p)

    return paragraphs

doc = Document()
preface, chapters = get_preface_and_chapters(home)
p = doc.add_paragraph(preface)
p.paragraph_format.first_line_indent = Pt(22)

for chapter in chapters:
    p = doc.add_heading(chapter.title, 1)

    for section in chapter.sections:
        p = doc.add_heading(section.title, 2)
        # p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraphs = get_paragraphs(section.url)
        for paragraph in paragraphs:
            if paragraph.img:
                doc.add_paragraph()

                try:
                    context = get(paragraph.img.url, None)

                    width, height = None, None
                    if paragraph.img.width:
                        width = Pt(paragraph.img.width)
                    if paragraph.img.height:
                        height = Pt(paragraph.img.height)
                    p = doc.add_picture(io.BytesIO(context), width=width, height=height)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except UnicodeDecodeError as e:
                    print("except:", e)
                except ZeroDivisionError as e:
                    print("except:", e)
                except LookupError as e:
                    print("except:", e)
                except UnrecognizedImageError as e:
                    print("except:", e)

                if paragraph.text:
                    p = doc.add_paragraph(paragraph.text)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

            elif paragraph.table:
                table = doc.add_table(rows=len(paragraph.table), cols=len(paragraph.table[0]))
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for r, row in enumerate(paragraph.table):
                    cells = table.rows[r].cells
                    for c, column in enumerate(row):
                        cells[c].text = column

            elif paragraph.text:
                p = doc.add_paragraph(paragraph.text)
                p.paragraph_format.first_line_indent = Pt(22)

doc.save("言必称希腊还是言必称中国.docx")

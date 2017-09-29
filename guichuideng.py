#!/usr/bin/env python3
# coding: utf-8

import requests
import lxml
import lxml.html
import lxml.etree
import parse
import chardet
import time
from docx import Document


urls = [
    ("鬼吹灯之精绝古城", "http://www.guichuideng.org/jing-jue-gu-cheng"),
    ("鬼吹灯之龙岭迷窟", "http://www.guichuideng.org/long-ling-mi-ku"),
    ("鬼吹灯之云南虫谷", "http://www.guichuideng.org/yun-nan-chong-gu"),
    ("鬼吹灯之昆仑神宫", "http://www.guichuideng.org/kun-lun-shen-gong"),
    ("鬼吹灯之黄皮子坟", "http://www.guichuideng.org/huang-pi-zi-fen"),
    ("鬼吹灯之南海归墟", "http://www.guichuideng.org/nan-hai-gui-xu"),
    ("鬼吹灯之怒晴湘西", "http://www.guichuideng.org/nu-qing-xiang-xi"),
    ("鬼吹灯之巫峡棺山", "http://www.guichuideng.org/wu-xia-guan-shan"),
]
encoding = "UTF-8"


session = requests.Session()
session.headers.update({
    "User-Agent": "Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; WOW64; Trident/5.0)",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Encoding": "gzip, deflate, sdch, br",
    "Accept-Language": "zh-CN,zh;q=0.8,en;q=0.6,zh-TW;q=0.4"
})


def get(url, encoding='utf-8', sleep_seconds=0):
    print("GET:", url)

    try:
        r = session.get(url, allow_redirects=True)
    except requests.exceptions.ConnectionError:
        if sleep_seconds == 0:
            sleep_seconds = 10
        time.sleep(sleep_seconds)

        return get(url, encoding, sleep_seconds + sleep_seconds)

    if r.status_code == 200:
        pass
    elif r.status_code == 404:
        raise LookupError("not found %s" % url)
    else:
        raise IOError('request "%s" got status: %d' % (url, r.status_code))

    if encoding is None:
        return r.content

    try:
        text = r.content.decode(encoding)
    except UnicodeDecodeError:
        encoding = chardet.detect(r.content)["encoding"]
        text = r.content.decode(encoding)

    return text


# chapter
def get_chapters(book_url, encoding='utf-8'):
    text = get(book_url, encoding)
    doc = lxml.html.fromstring(text)

    body = doc.xpath("/html/body")[0]
    articles = body.xpath(".//article[@class=\"excerpt excerpt-c3\"]")

    chapters = []
    for article in articles:
        a = article.xpath("a")[0]

        title = parse.parse("{} {}", a.text)[1]
        link = a.attrib["href"]
        chapters.append((title, link), )

    return chapters


def get_content(chapter_url, encoding="utf-8"):
    text = get(chapter_url, encoding)
    doc = lxml.html.fromstring(text)

    body = doc.xpath("/html/body")[0]
    lines = body.xpath(".//article[@class=\"article-content\"]/p")

    return list(map(lambda line: line.text, lines))

for url in urls:
    book_title = url[0]
    book_url = url[1]

    chapters = get_chapters(book_url, encoding)

    doc = Document()

    for chapter in chapters:

        doc.add_heading(chapter[0])

        lines = get_content(chapter[1], encoding)
        for line in lines:
            doc.add_paragraph(line)

    doc.save("{}.docx".format(book_title))
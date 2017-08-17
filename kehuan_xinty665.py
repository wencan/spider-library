#!/usr/bin/env python3
# coding: utf-8

import requests
import lxml
import lxml.html
import lxml.etree
import urllib.parse
import parse
import os.path
import io
from docx import Document

home = "http://www.readers365.com/kehuan/"
encoding = "GB18030"


def get(url, encoding='utf-8', retry=True):
    session = requests.Session()
    session.headers.update({
        "User-Agent": "Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; WOW64; Trident/5.0)",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Encoding": "gzip, deflate, sdch, br",
        "Accept-Language": "zh-CN,zh;q=0.8,en;q=0.6,zh-TW;q=0.4"
    })

    print("GET:", url)
    r = session.get(url, allow_redirects=True)
    if r.status_code == 200:
        pass
    elif r.status_code == 404:
        if retry:
            return get(url, encoding, False)
        else:
            raise LookupError("not found %s" % url)
    else:
        raise IOError('request "%s" got status: %d' % (url, r.status_code))

    if encoding is None:
        return r.content

    text = r.content.decode(encoding)
    args = parse.parse('''{}self.location = "{}";{}''', text)
    if args:
        target = args[1]
        url = urllib.parse.urljoin(url, target)

        return get(url, encoding)
    else:
        return text


# home
def get_books(home_url, encoding='utf-8'):
    text = get(home_url, encoding)
    doc = lxml.html.fromstring(text)

    body = doc.xpath("/html/body")[0]
    table = body.xpath("//table[@id='table2']")[0]

    links = table.xpath("//tr/td/a/@href")
    links = list(map(lambda link: "{}/index.htm".format(link.split("=")[0]) if "=" in link else link, links))
    links = list(map(lambda link: link.replace("00.htm", "index.htm"), links))
    links = list(map(lambda link: urllib.parse.urljoin(home_url, link), links))
    imgs = table.xpath("//tr/td/a/img/@src")
    imgs = list(map(lambda img: urllib.parse.urljoin(home_url, img), imgs))
    titles = table.xpath("//tr/td/font")
    titles = list(map(lambda title: title.text, titles))

    books = list(map(lambda idx: (titles[idx], links[idx], imgs[idx]), range(len(titles))))
    return books


# chapter
def get_chapters(book_url, encoding='utf-8'):
    text = get(book_url, encoding)
    doc = lxml.html.fromstring(text)

    body = doc.xpath("/html/body")[0]
    divs = body.xpath(".//div[@class=\"TitleLinks\"]")

    chapters = []
    if divs:
        div = divs[0]
        for line in div.iter():
            title = None
            href = None
            header = False

            if line.tag == "a":
                title = line.text
                href = urllib.parse.urljoin(book_url, line.attrib["href"])
                if not title:
                    fonts = line.xpath("font[@size=\"4\"]")
                    if fonts:
                        title = fonts[0].text
                        if fonts[0].tail:
                            title = "{} {}".format(title, fonts[0].tail)
                        header = True
                if not title:
                    continue
            elif line.tag == "font":
                parent = line.getparent()
                if parent is not None and parent.tag == "a":
                    continue
                title = line.text
                header = True
                if not title:
                    continue
            else:
                continue

            title = title.rstrip()
            title = title.lstrip()
            title = title.replace("\u3000", "", -1)
            chapters.append((title, href, header),)
    else:
        table = body.xpath("//td[@width=\"100%\"]/div[@align=\"center\"]/table[@border=\"0\"]")[0]
        hrefs = table.xpath(".//a[@href]")
        for href in hrefs:
            title = href.text
            link = urllib.parse.urljoin(book_link, href.attrib["href"])
            chapters.append((title, link, False), )
    return chapters


def get_content(chapter_url, encoding="utf-8"):
    text = get(chapter_url, encoding)
    doc = lxml.html.fromstring(text)

    try:
        body = doc.xpath("/html/body")[0]
        divs = body.xpath(".//div[@align=\"center\"]")

        center = None
        for div in divs:
            header = None
            tailer = None
            for node in div.iter():
                if isinstance(node, lxml.html.HtmlComment):
                    if node.text == "HTMLBUILERPART0":
                        header = node.getparent()
                        continue
                    elif node.text == "/HTMLBUILERPART0":
                        tailer = node.getparent()
            if header is not None and tailer is not None:
                centers = list(filter(lambda node: header in node.iter() and tailer in node.iter(), div.iter()))
                center = centers[len(centers)-1]
                break
        if center is None:
            return None

        context = []
        for node in center.iter():
            if node.tag == "br":
                if node.tail:
                    context.append((node.tail, False))
            elif node.tag == "img":
                if "src" in node.attrib:
                    src = urllib.parse.urljoin(chapter_url, node.attrib["src"])
                    context.append((src, True))
            elif node.tag == "div":
                if node.tail:
                    context.append((node.tail, False))

        return context
    except IndexError:
        print(text)
        raise


books = get_books(home, encoding)
for idx, book in enumerate(books):
    book_title = book[0]
    book_link = book[1]
    book_img = book[2]

    print(idx, book_title)

    _, suffix = os.path.splitext(book_img)
    try:
        buff = get(book_img, None)
    except LookupError as e:
        print(e)
    else:
        book_img = "{}{}".format(book_title, suffix)
        with open(book_img, mode="w+b") as f:
            f.write(buff)

    doc = Document()

    chapters = get_chapters(book_link, encoding)
    for idx, chapter in enumerate(chapters):
        chapter_title = chapter[0]
        chapter_link = chapter[1]
        chapter_is_header = chapter[2]

        if idx != 0 and not (chapters[idx-1][1] is None and chapters[idx-1][2] is True):
            doc.add_page_break()
        doc.add_heading(chapter_title, level=chapter_is_header and 2 or 3)

        if not chapter_link:
            continue

        lines = get_content(chapter_link, encoding)
        if not lines:
            print("blank page")
            continue
        for line in lines:
            text, is_img = line[0], line[1]
            if is_img:
                context = get(text, None)
                try:
                    doc.add_picture(io.BytesIO(context))
                except UnicodeDecodeError as e:
                    print(e)
            else:
                text = text.strip("\r\n ")
                doc.add_paragraph(text)

    doc.save("{}.docx".format(book_title))
